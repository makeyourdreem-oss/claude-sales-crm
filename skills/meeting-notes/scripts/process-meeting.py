"""
Превращает расшифровку встречи в 3-секционный DOCX:
1. Транскрибация (без изменений)
2. Саммари
3. Выводы и Action Items

Саммари и conclusions передаются как отдельные текстовые файлы — их генерирует
LLM в основной сессии, а этот скрипт собирает финальный документ.

Usage:
    python process-meeting.py <input.docx|.txt> --summary <s.txt> --conclusions <c.txt> [--output <p>] [--style <name>]
"""
import sys
import argparse
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print('ERROR: pip install python-docx', file=sys.stderr)
    sys.exit(1)


DEFAULT_STYLE = {
    'title_color': '#1A1A2E',
    'h2_color': '#16548E',
    'body_color': '#2D2D2D',
    'footer_color': '#666666',
    'title_size': 18,
    'h2_size': 13,
    'body_size': 11,
}


def hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def read_input(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == '.docx':
        doc = Document(str(path))
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    if suffix in ('.txt', '.md'):
        return path.read_text(encoding='utf-8')
    raise ValueError(f'Unsupported input: {suffix}')


def load_style(name: str | None) -> dict:
    if not name or name == 'default':
        return DEFAULT_STYLE
    styles_dir = Path(__file__).parent.parent / 'styles'
    style_file = styles_dir / f'{name}.json'
    if style_file.exists():
        return {**DEFAULT_STYLE, **json.loads(style_file.read_text(encoding='utf-8'))}
    print(f'WARN: style "{name}" not found, using default', file=sys.stderr)
    return DEFAULT_STYLE


def style_run(run, size, color, bold=False):
    run.font.size = Pt(size)
    run.font.color.rgb = hex_to_rgb(color)
    run.font.bold = bold


def add_section(doc, title, body, style):
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        style_run(run, style['h2_size'], style['h2_color'], bold=True)
    for chunk in body.split('\n\n'):
        if chunk.strip():
            p = doc.add_paragraph(chunk.strip())
            for run in p.runs:
                style_run(run, style['body_size'], style['body_color'])


def build_doc(transcript, summary, conclusions, output, title, style):
    doc = Document()
    t = doc.add_heading(title, level=1)
    for run in t.runs:
        style_run(run, style['title_size'], style['title_color'], bold=True)
    meta = doc.add_paragraph(f'Обработано: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    for run in meta.runs:
        style_run(run, 10, style['footer_color'])
    add_section(doc, '1. Транскрибация (без изменений)', transcript, style)
    add_section(doc, '2. Саммари', summary, style)
    add_section(doc, '3. Выводы и Action Items', conclusions, style)
    doc.save(str(output))
    return output


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('input')
    ap.add_argument('--summary', required=True)
    ap.add_argument('--conclusions', required=True)
    ap.add_argument('--output')
    ap.add_argument('--title', default='Конспект встречи')
    ap.add_argument('--style', default='default')
    args = ap.parse_args()

    input_path = Path(args.input)
    transcript = read_input(input_path)
    summary = Path(args.summary).read_text(encoding='utf-8')
    conclusions = Path(args.conclusions).read_text(encoding='utf-8')

    output = Path(args.output) if args.output else input_path.parent / f'Конспект_{input_path.stem}_{datetime.now().strftime("%Y-%m-%d")}.docx'

    style = load_style(args.style)
    build_doc(transcript, summary, conclusions, output, args.title, style)
    print(f'OK: {output}')


if __name__ == '__main__':
    main()
